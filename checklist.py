import docx
import zipfile
from bs4 import BeautifulSoup

doc = (r"D:\Users\guima\OneDrive\UFSC\Programação\PMRR\Exemplos\P 001 BURACO NA VIA PÚBLICA  02 01 2023.docx")

bestand = docx.Document(doc)
tabellen = bestand.tables

#get data from all the "normal" fields

alletabellen = []     
for i, tabel in enumerate(tabellen):
    for row in tabellen[i].rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                alletabellen.append(paragraph.text)

#get data from all the dropdown lists

document = zipfile.ZipFile(doc)
xml_data = document.read('word/document.xml')
document.close()

soup = BeautifulSoup(xml_data, 'xml')
gegevens = soup.findAll('ddList')     #search dropdownlists (n = 12)

dropdownlist = []
dropdownlistdata = []

for i in gegevens:
    dropdownlist.append(i.find('result'))

#convert to string for if statements
number = str(dropdownlist[0])
job = str(dropdownlist[1])
vehicle = str(dropdownlist[7])

if number == '<w:result w:val="1"/>' :
    dropdownlistdata.append('0,3')
elif number == '<w:result w:val="2"/>' :
    dropdownlistdata.append('1,2')
elif number == '<w:result w:val="3"/>' :
    dropdownlistdata.append('onbekend')
else:
    dropdownlistdata.append('geen')

if job  == '<w:result w:val="1"/>' :
    dropdownlistdata.append('nee')
else:
    dropdownlistdata.append('ja')

if vehicle == '<w:result w:val="1"/>' :
    dropdownlistdata.append('nee')
else:
    dropdownlistdata.append('ja')

#show data
print (alletabellen)
print (dropdownlistdata)