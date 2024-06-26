
from docx import Document
import pandas as pd
import zipfile

# Load the DOC file using double quotes
doc = Document(r"D:\Users\guima\OneDrive\UFSC\Programação\PMRR\Exemplos\P 001 BURACO NA VIA PÚBLICA  02 01 2023.docx")

form_data = []
table_data = []
checklist_data = []

# Extract text from paragraphs (handling multi-line fields and checklists)
current_field = None
for para in doc.paragraphs:
    text = para.text.strip()
    
    if ":" in text:
        if current_field:
            form_data.append(current_field)
        key, value = text.split(":", 1)
        current_field = [key.strip(), value.strip()]
    elif text.startswith(u"\\u2022") or text.startswith(u"\\u25A0"):
        # Identify bullet points or checkboxes
        checklist_data.append(text)
    else:
        if current_field:
            current_field[1] += " " + text

if current_field:
    form_data.append(current_field)

# Extract text from tables (handling nested tables if any)
def extract_table_data(table):
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
        for cell in row.cells:
            if cell.tables:
                for nested_table in cell.tables:
                    extract_table_data(nested_table)

for table in doc.tables:
    extract_table_data(table)

# Convert form data to DataFrame
df_form = pd.DataFrame(form_data, columns=["Field", "Value"])

# Convert table data to DataFrame (if all tables have the same structure)
df_table = pd.DataFrame(table_data)

# Convert checklist data to DataFrame
df_checklist = pd.DataFrame(checklist_data, columns=["Checklist Item"])

# Save form data, table data, and checklist data to separate sheets in an Excel file
with pd.ExcelWriter("output.xlsx") as writer:
    df_form.to_excel(writer, sheet_name="Form Data", index=False)
    df_table.to_excel(writer, sheet_name="Table Data", index=False)
    df_checklist.to_excel(writer, sheet_name="Checklist Data", index=False)